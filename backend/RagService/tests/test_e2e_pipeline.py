"""End-to-end offline test of the whole ingestion pipeline.

Exercises the real orchestration (IngestionService driven by the real IngestionWorker)
wired to FAKE providers only — FakeFileStorage feeds in-code fixture bytes,
FakeEmbeddingProvider returns EMBEDDING_DIM vectors, and the repositories are
in-memory. No Ollama, no Postgres. The extractor/OCR/chunker are the real ones.
"""

from __future__ import annotations

from uuid import uuid4

import pytest

from app.application.services.ingestion_service import IngestionJob
from app.application.services.ingestion_worker import IngestionWorker
from app.domain.entities.document import Document
from app.domain.enums.chunk_source import ChunkSource
from app.domain.enums.document_status import DocumentStatus
from app.infrastructure.config.settings import Settings
from app.infrastructure.ocr.tesseract_ocr_processor import tesseract_available

from tests.extraction.fixtures import DOCX_CONTENT_TYPE, make_docx
from tests.ingestion.fakes import (
    FakeClock,
    FakeEmbeddingProvider,
    FakeFileStorage,
    InMemoryChunkRepository,
    InMemoryDocumentRepository,
    build_ingestion_service,
)

DIM = Settings().embedding_dim
PDF_CONTENT_TYPE = "application/pdf"


def _seed_pending(documents: InMemoryDocumentRepository, job: IngestionJob) -> None:
    documents.seed(
        Document(
            classroom_id=job.classroom_id, file_id=job.file_id, s3_key=job.s3_key,
            file_name=job.file_name, content_type=job.content_type,
            status=DocumentStatus.PENDING,
        )
    )


async def _process(service, *jobs: IngestionJob) -> None:
    """Run jobs through a real worker to completion (no retries here, so join()
    reliably drains)."""
    worker = IngestionWorker(
        lambda job: service.ingest(job), concurrency=1, queue_max=32, clock=FakeClock()
    )
    await worker.start()
    try:
        for job in jobs:
            worker.enqueue(job)
        await worker.join()
    finally:
        await worker.stop()


async def test_e2e_docx_full_pipeline_via_worker() -> None:
    job = IngestionJob(uuid4(), uuid4(), "class/doc.docx", "doc.docx", DOCX_CONTENT_TYPE)
    documents = InMemoryDocumentRepository()
    chunks = InMemoryChunkRepository()
    embedder = FakeEmbeddingProvider(DIM)
    service = build_ingestion_service(
        storage=FakeFileStorage({job.s3_key: make_docx()}),
        embedder=embedder, documents=documents, chunks=chunks, clock=FakeClock(),
    )
    _seed_pending(documents, job)

    await _process(service, job)

    stored = await documents.get_by_file_id(job.file_id)
    assert stored.status == DocumentStatus.DONE
    assert stored.content_hash is not None
    assert stored.attempts == 1

    persisted = chunks.by_document[stored.id]
    assert len(persisted) >= 1
    assert [chunk.chunk_index for chunk, _ in persisted] == list(range(len(persisted)))
    for _, embedding in persisted:
        assert len(embedding) == DIM
    # docx heading structure is preserved as section metadata.
    assert any("section" in chunk.metadata for chunk, _ in persisted)
    assert embedder.embed_documents_calls >= 1


async def test_e2e_reindex_replaces_chunks_on_changed_content() -> None:
    job = IngestionJob(uuid4(), uuid4(), "class/doc.docx", "doc.docx", DOCX_CONTENT_TYPE)
    documents = InMemoryDocumentRepository()
    chunks = InMemoryChunkRepository()
    storage = FakeFileStorage({job.s3_key: make_docx()})
    service = build_ingestion_service(
        storage=storage, embedder=FakeEmbeddingProvider(DIM),
        documents=documents, chunks=chunks, clock=FakeClock(),
    )
    _seed_pending(documents, job)
    await _process(service, job)

    stored = await documents.get_by_file_id(job.file_id)
    first_ids = {chunk.id for chunk, _ in chunks.by_document[stored.id]}

    # Re-upload changed content, then reindex (reset to Pending as the endpoint does).
    storage.put(job.s3_key, _other_docx())
    await documents.reset_to_pending(job.file_id, reset_attempts=True)
    await _process(service, job)

    reprocessed = await documents.get_by_file_id(job.file_id)
    assert reprocessed.status == DocumentStatus.DONE
    assert reprocessed.attempts == 1  # reset then one fresh claim
    new_ids = {chunk.id for chunk, _ in chunks.by_document[reprocessed.id]}
    assert new_ids.isdisjoint(first_ids)  # a brand-new chunk set replaced the old one


def _other_docx() -> bytes:
    import io

    from docx import Document as open_docx

    document = open_docx()
    document.add_heading("Different", level=1)
    document.add_paragraph("Totally different content for the reindex round trip.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def test_e2e_delete_removes_document_and_chunks() -> None:
    job = IngestionJob(uuid4(), uuid4(), "class/doc.docx", "doc.docx", DOCX_CONTENT_TYPE)
    documents = InMemoryDocumentRepository()
    chunks = InMemoryChunkRepository()
    service = build_ingestion_service(
        storage=FakeFileStorage({job.s3_key: make_docx()}),
        embedder=FakeEmbeddingProvider(DIM), documents=documents, chunks=chunks,
        clock=FakeClock(),
    )
    _seed_pending(documents, job)
    await _process(service, job)
    stored = await documents.get_by_file_id(job.file_id)

    await chunks.delete_by_document_id(stored.id)
    await documents.delete_by_file_id(job.file_id)

    assert await documents.get_by_file_id(job.file_id) is None
    assert stored.id not in chunks.by_document


async def test_e2e_forced_failure_marks_failed_and_worker_survives() -> None:
    documents = InMemoryDocumentRepository()
    chunks = InMemoryChunkRepository()
    bad = IngestionJob(uuid4(), uuid4(), "class/bad.pdf", "bad.pdf", PDF_CONTENT_TYPE)
    good = IngestionJob(uuid4(), uuid4(), "class/good.docx", "good.docx", DOCX_CONTENT_TYPE)
    service = build_ingestion_service(
        storage=FakeFileStorage(
            {bad.s3_key: b"%PDF-1.4 not a real pdf", good.s3_key: make_docx()}
        ),
        embedder=FakeEmbeddingProvider(DIM), documents=documents, chunks=chunks,
        clock=FakeClock(),
    )
    _seed_pending(documents, bad)
    _seed_pending(documents, good)

    await _process(service, bad, good)

    assert (await documents.get_by_file_id(bad.file_id)).status == DocumentStatus.FAILED
    # The bad document did not kill the worker; the good one still completed.
    good_doc = await documents.get_by_file_id(good.file_id)
    assert good_doc.status == DocumentStatus.DONE
    assert len(chunks.by_document[good_doc.id]) >= 1


@pytest.mark.skipif(
    not tesseract_available(),
    reason="tesseract binary not installed — OCR recovery is skipped",
)
async def test_e2e_scanned_pdf_recovers_text_via_ocr() -> None:
    from tests.ocr.fixtures import make_scanned_pdf, normalize

    job = IngestionJob(uuid4(), uuid4(), "class/scan.pdf", "scan.pdf", PDF_CONTENT_TYPE)
    documents = InMemoryDocumentRepository()
    chunks = InMemoryChunkRepository()
    service = build_ingestion_service(
        storage=FakeFileStorage({job.s3_key: make_scanned_pdf()}),
        embedder=FakeEmbeddingProvider(DIM), documents=documents, chunks=chunks,
        clock=FakeClock(),
    )
    _seed_pending(documents, job)

    await _process(service, job)

    stored = await documents.get_by_file_id(job.file_id)
    assert stored.status == DocumentStatus.DONE
    persisted = chunks.by_document[stored.id]
    assert any(chunk.source == ChunkSource.OCR for chunk, _ in persisted)
    recovered = normalize(" ".join(chunk.text for chunk, _ in persisted))
    assert "quick brown fox" in recovered
